#!/usr/bin/env python3
"""
Generate the comprehensive Watershed Capital Project Handover Document.
14 sections covering everything from architecture to sign-off.
"""
import sys, os
DOCX_SCRIPTS = "/home/z/my-project/skills/docx/scripts"
if DOCX_SCRIPTS not in sys.path:
    sys.path.insert(0, DOCX_SCRIPTS)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ============================================================================
# COLOR PALETTE — Watershed Capital brand (emerald/teal)
# ============================================================================
COLOR_PRIMARY = RGBColor(0x04, 0x78, 0x57)      # Emerald 700
COLOR_SECONDARY = RGBColor(0x0D, 0x94, 0x88)     # Teal 600
COLOR_ACCENT = RGBColor(0x1E, 0x40, 0xAF)        # Blue 800
COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)          # Slate 900
COLOR_GRAY = RGBColor(0x64, 0x74, 0x8B)          # Slate 500
COLOR_LIGHT_BG = "F0FDF4"                          # Emerald 50
COLOR_TABLE_HEADER = "047857"                      # Emerald 700 hex
COLOR_TABLE_ALT = "F0FDF4"                          # Emerald 50 hex

doc = Document()

# ============================================================================
# DOCUMENT SETUP — Page margins, default fonts
# ============================================================================
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.3

# Heading styles
for i, (size, color) in enumerate([(18, COLOR_PRIMARY), (15, COLOR_SECONDARY), (13, COLOR_ACCENT), (12, COLOR_DARK)]):
    h = doc.styles[f'Heading {i+1}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if i == 0 else 14)
    h.paragraph_format.space_after = Pt(8)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def add_cover():
    """Professional cover page using a full-page table."""
    # Outer table for background
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(16)
    # Set background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="047857" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    # Set row height
    tr = tbl.rows[0]._tr
    trPr = parse_xml(f'<w:trPr {nsdecls("w")}><w:trHeight w:val="9500" w:hRule="exact"/></w:trPr>')
    tr.append(trPr)

    # Content
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("WATERSHED CAPITAL")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Banking Governance Platform")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0xA7, 0xF3, 0xD0)
    run2.font.name = 'Calibri'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(40)
    run3 = p3.add_run("PROJECT HANDOVER DOCUMENT")
    run3.font.size = Pt(22)
    run3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run3.font.bold = True
    run3.font.name = 'Calibri'

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Comprehensive System Documentation")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0xD1, 0xFA, 0xE5)
    run4.italic = True

    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(80)
    run5 = p5.add_run(f"Version 14.0  |  {datetime.date.today().strftime('%B %d, %Y')}")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0xA7, 0xF3, 0xD0)

    p6 = cell.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("Prepared by: Development Team\nFor: Watershed Capital Management")
    run6.font.size = Pt(10)
    run6.font.color.rgb = RGBColor(0xD1, 0xFA, 0xE5)

    doc.add_page_break()


def add_heading(text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, bold=False, italic=False, size=None, color=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.3
    p.add_run(text)
    return p


def add_table(headers, rows, col_widths=None):
    """Add a formatted table with header row and alternating row colors."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            # Alternating row color
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_ALT}" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_code_block(text):
    """Add a monospace code block with dark background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_DARK
    # Light gray background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    # Border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/>'
        f'</w:pBdr>'
    )
    p._p.get_or_add_pPr().append(pBdr)


def add_callout(title, text, color="047857"):
    """Add a highlighted callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0FDF4" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    run = p.add_run(f"⚠ {title}: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

# ============================================================================
# BUILD DOCUMENT
# ============================================================================

add_cover()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. System Architecture Document",
    "3. Installation & Setup Guide",
    "4. Deployment Guide",
    "5. Environment Variables Template",
    "6. Database Documentation (including ERD)",
    "7. API Documentation",
    "8. Administrator Manual",
    "9. End-User Manual",
    "10. Security & Access Control Guide",
    "11. Backup & Disaster Recovery Guide",
    "12. Maintenance & Troubleshooting Guide",
    "13. Source Code Repository Information",
    "14. Credentials & Asset Handover Checklist",
    "15. Project Acceptance & Sign-off Document",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.size = Pt(11)

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
add_heading("1. Executive Summary", level=1, page_break=True)

add_heading("1.1 Project Overview", level=2)
add_para("The Watershed Capital Banking Governance Platform is a comprehensive, bank-grade Loan Origination System (LOS) and Credit Appraisal Memorandum (CAM) platform designed for Microfinance Banks, Finance Companies, Cooperative Banks, and Commercial Banks operating under Central Bank of Nigeria (CBN) regulatory guidelines. The system manages the complete credit lifecycle from customer onboarding through loan origination, multi-level approval, disbursement, post-disbursement monitoring, and loan closure.")

add_para("The platform implements a 13-step bank-grade loan origination workflow with 5 phases: Pre-Qualification, Structuring (Engine Room), Governance, Closing, and Post-Disbursement Monitoring. Each department (LO, Legal, BM, HOC, Analyst, CRO, CFO, MD) owns its own recommendation — no department overwrites another's opinion, ensuring full accountability and audit trail.")

add_heading("1.2 Key Metrics", level=2)
add_table(
    ["Metric", "Value"],
    [
        ["Technology Stack", "Next.js 16 (Turbopack), React 19, TypeScript, Prisma ORM, SQLite/Turso"],
        ["Source Files", "294 TypeScript/TSX files, 41,722 lines of code"],
        ["API Routes", "137 RESTful API endpoints"],
        ["UI Views", "74 component views"],
        ["Database Models", "79 Prisma models"],
        ["CAM Tabs", "12-tab Universal CAM with 30+ financial formulas"],
        ["Workflow Steps", "13-step origination + 5-step post-disbursement monitoring"],
        ["Approval Levels", "8-level MCC (Management Credit Committee) chain"],
        ["Roles", "13 staff roles with 60+ permission flags"],
        ["Sectors Seeded", "88 business sectors with benchmarked margins"],
        ["Locations", "60+ Nigerian locations with risk ratings"],
        ["Authentication", "JWT-based (HMAC-SHA256, 8-hour expiry)"],
        ["Security", "CBN-compliant: DSR ≤45%, Gearing ≤35%, Collateral ≥100%"],
    ],
    col_widths=[5, 11]
)

add_heading("1.3 Core Modules", level=2)
modules = [
    ("Credit Governance", "13-step loan origination workflow, Universal CAM with forensic financial analysis, 8-snapshot audit trail"),
    ("MCC Committee", "8-level Management Credit Committee decision ledger with CP checklist (22 mandatory documents)"),
    ("Treasury", "Investor onboarding, product configuration, deal booking, redemptions, bank assets, profitability reporting"),
    ("Accounting & GL", "Double-entry accounting, chart of accounts, journal entries, financial statements, bank reconciliation, invoicing, expenses, payroll, teller operations, tills, AP/AR"),
    ("Governance & Compliance", "Audit trail, login history, activity log, compliance monitoring, policy documents, conditions precedent, pre-disbursement checklist, risk assessment, exception reports"),
    ("Communication", "Announcements, message center, notification center, 12 email templates, SMS broadcast, 7 drip campaigns, customer service (FAQ, chat, callback, restructuring, receipts)"),
    ("Customer Portal", "Self-service portal with loan application, offers, repayment, savings, investments, transactions, documents, KYC, bank accounts, security, referral, FAQ, chat"),
    ("System Administration", "Branch network (6 branches), staff & access control (13 roles), loan products, business sectors, branding, global configuration"),
]
for name, desc in modules:
    add_bullet(f"{name}: {desc}")

add_heading("1.4 Compliance & Standards", level=2)
add_para("The system enforces CBN prudential guidelines including:")
add_bullet("Debt Service Ratio (DSR) hard gate at 45% — automatic submission block")
add_bullet("Gearing Ratio hard gate at 35% — automatic submission block")
add_bullet("Collateral coverage minimum 100% FSV — critical red flag if below")
add_bullet("Single obligor limit at ₦50M (configurable) — blocks submission if exceeded")
add_bullet("12-month solvency projection — blocks submission if insolvent")
add_bullet("Engine REJECT verdict — blocks CAM submission without MD/CEO override")
add_bullet("AML/CFT thresholds: CTR at ₦5M (individual) / ₦10M (corporate)")
add_bullet("Snapshot locking — once locked, CAM data is immutable (only super-admin/MD can override with documented reason)")
add_bullet("Full audit trail — every CAM edit, transition, and override is logged")

# ============================================================================
# 2. SYSTEM ARCHITECTURE
# ============================================================================
add_heading("2. System Architecture Document", level=1, page_break=True)

add_heading("2.1 Architecture Overview", level=2)
add_para("The platform uses a modern full-stack JavaScript/TypeScript architecture built on Next.js 16 with the App Router pattern. The frontend and backend are co-located in a single Next.js application, with API routes serving as the backend layer and React components as the frontend.")

add_heading("2.2 Technology Stack", level=2)
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Frontend", "Next.js 16, React 19, TypeScript", "Server-side rendering, client-side hydration, type safety"],
        ["UI Framework", "Tailwind CSS 4, shadcn/ui", "Utility-first styling, accessible component library"],
        ["State Management", "Zustand (with persist middleware)", "Client-side state, localStorage session persistence"],
        ["Backend", "Next.js API Routes (137 endpoints)", "RESTful API, server-side business logic"],
        ["Database", "SQLite (dev), Turso/PostgreSQL (prod)", "Prisma ORM with type-safe queries"],
        ["ORM", "Prisma 6.11", "Schema-first database modeling, migrations, type-safe client"],
        ["Authentication", "JWT (HMAC-SHA256 via Node.js crypto)", "Stateless token-based auth, 8-hour expiry"],
        ["PDF Generation", "@react-pdf/renderer", "CAM memo, offer letters, repayment schedules"],
        ["Charts", "Recharts", "Dashboard visualizations"],
        ["Icons", "Lucide React", "Consistent iconography"],
        ["Package Manager", "Bun (recommended) / npm", "Dependency management, script runner"],
    ],
    col_widths=[3.5, 5, 7.5]
)

add_heading("2.3 Application Structure", level=2)
add_code_block("""src/
├── app/
│   ├── api/                    # 137 API route handlers
│   │   ├── auth/               # Login, logout
│   │   ├── loans/              # Loan CRUD, transitions, disbursement
│   │   ├── appraisals/         # CAM save/load
│   │   ├── engine/             # Credit engine recalculation
│   │   ├── mcc/                # MCC decisions, checklist
│   │   ├── treasury/           # Treasury operations
│   │   ├── accounting/         # GL, journal, invoices, expenses
│   │   ├── compliance/         # Conditions, documents, verification
│   │   ├── communications/     # Announcements, messages, notifications
│   │   ├── customers/          # Customer database
│   │   ├── sectors/            # 88 business sectors
│   │   ├── branches/           # Branch network
│   │   ├── staff/              # Staff management
│   │   └── ... (40+ more)
│   ├── page.tsx                # Main view router (admin/customer/public)
│   ├── layout.tsx              # Root layout
│   └── globals.css             # Global styles
├── components/
│   ├── views/                  # 74 UI views
│   │   ├── cam.tsx             # Universal CAM (4,200+ lines)
│   │   ├── dashboard.tsx       # Admin dashboard
│   │   ├── loan/               # Loan list & detail
│   │   ├── mcc/                # MCC decision center
│   │   ├── treasury/           # Treasury module
│   │   ├── accounting/         # Accounting module
│   │   ├── governance/         # Audit & compliance
│   │   ├── system/             # Branch, staff, sectors, settings
│   │   ├── communications/     # 7 communication views
│   │   ├── customer/           # Customer portal (20+ views)
│   │   └── public/             # Public website
│   ├── ui/                     # shadcn/ui components
│   ├── pdf/                    # PDF document templates
│   ├── sidebar.tsx             # Navigation sidebar
│   └── topbar.tsx              # Top navigation bar
├── lib/
│   ├── credit-engine.ts        # 30+ financial formulas (2,600 lines)
│   ├── auth.ts                 # JWT authentication (server-side)
│   ├── auth-client.ts          # Token management (client-side)
│   ├── constants.ts            # Workflow, roles, permissions, sectors
│   ├── db.ts                   # Prisma client singleton
│   ├── store.ts                # Zustand state management
│   ├── loan-calc.ts            # Loan schedule calculator
│   ├── notifications.ts        # Notification system
│   └── email-service.ts        # Multi-provider email service
├── prisma/
│   └── schema.prisma           # 79 models, 2,240 lines
└── scripts/
    ├── seed.ts                 # Infrastructure data seeding
    └── check-db.ts             # Database verification""")

add_heading("2.4 Loan Origination Workflow (13 Steps)", level=2)
add_para("The system implements a bank-grade 13-step loan origination workflow across 5 phases:")
add_table(
    ["Step", "Phase", "Actor", "Action"],
    [
        ["1", "Pre-Qualification", "Loan Officer (LO)", "Data collection, KYC upload"],
        ["2", "Pre-Qualification", "Legal Team", "CAC verification — early gatekeeper (reject if inactive)"],
        ["3", "Pre-Qualification", "Branch Manager (BM)", "Field & document verification"],
        ["4", "Structuring", "Head of Credit (HOC)", "Assigns best analyst by sector/loan size"],
        ["5", "Structuring", "Credit Analyst", "Forensic audit, CAM, Triple Lock (Amount/Tenor/Rate)"],
        ["6", "Structuring", "HOC", "Quality assurance — validates analyst's math"],
        ["7", "Governance", "CRO", "Risk audit — provides max safe exposure (opinion)"],
        ["8", "Governance", "CFO", "Liquidity check — provides liquidity limit (opinion)"],
        ["9", "Governance", "Legal", "Aggregation — compiles Executive Credit Pack (sole aggregator)"],
        ["10", "Governance", "MD/CEO", "Final approval — auto-recalculates schedule if amount changes"],
        ["11", "Closing", "Customer", "Accepts or negotiates offer"],
        ["12", "Closing", "HOC", "Go-Live — activates loan account (status = RUNNING)"],
        ["13", "Closing", "CFO/Treasury", "Disbursement — releases funds (pre-disbursement validation)"],
    ],
    col_widths=[1.2, 3, 3.5, 8.3]
)

add_heading("2.5 Post-Disbursement Monitoring", level=2)
add_para("After disbursement, loans enter a 5-state monitoring phase:")
add_bullet("Active Monitoring — ongoing portfolio health checks")
add_bullet("Repayment Tracking — due date tracking, payment matching")
add_bullet("Early Warning Signals — NPL classification (8 CBN statuses)")
add_bullet("Collections — delinquent loan recovery")
add_bullet("Loan Closure — full repayment or write-off")

add_heading("2.6 Credit Engine", level=2)
add_para("The credit engine (src/lib/credit-engine.ts) implements 30+ financial formulas with CBN hard gates:")
add_bullet("Sales Forensics: 4-source triangulation (Client Estimate, Spot Check, Bank Statement, Book Records) — picks least figure")
add_bullet("Weighted Margin: Cost-weighted margin calculation with sector benchmark comparison")
add_bullet("Purchase Verification: Derived COGS = Sales × (1 − Weighted Margin)")
add_bullet("PMT: Standard amortization (Flat & Reducing Balance)")
add_bullet("Ratios: DSR, DSCR, Gearing, Current Ratio, Quick Ratio, DIO, DSO, DPO, Cash Conversion Cycle")
add_bullet("12-Month Projection: Monthly cashflow with insolvency detection")
add_bullet("Stress Test: Sales haircut, margin compression, OPEX increase simulation")
add_bullet("Risk Grade: A/B/C/D/F with APPROVE/REVIEW/REJECT verdicts")
add_bullet("Red Flags: 8 flags with severity levels and point deductions")
add_bullet("Collateral Coverage: FSV haircuts (Movable 80%, Immovable 60%, Cash 100%, Stock 10%)")
add_bullet("Bank Yield: Interest income, processing fee, CCD income, cost of fund (30%), admin cost (5%), net yield")
add_bullet("Detailed Cashflow: 22-row × 12-month Excel parity table")
add_bullet("Cost-of-Fund Schedule: 30% PA amortization")
add_bullet("Convert-to-Loan Schedule: Principal + upfront + CCD + admin amortized")

# ============================================================================
# 3. INSTALLATION & SETUP GUIDE
# ============================================================================
add_heading("3. Installation & Setup Guide", level=1, page_break=True)

add_heading("3.1 Prerequisites", level=2)
add_bullet("Node.js 18+ (recommended: Node.js 20 LTS)")
add_bullet("Bun 1.0+ (recommended package manager — faster, more resilient)")
add_bullet("SQLite (included with Node.js) or PostgreSQL 14+ for production")
add_bullet("Git for version control")

add_heading("3.2 Installation Steps", level=2)
add_para("Step 1: Extract the project archive", bold=True)
add_code_block("unzip watershed-capital-app-v14.zip\ncd watershed-capital-app-v14")

add_para("Step 2: Install dependencies (use Bun for best results)", bold=True)
add_code_block("# Option A: Bun (recommended — faster, no network errors)\nbun install\n\n# Option B: npm (slower, may need retry)\nnpm install")

add_para("Step 3: Configure environment variables", bold=True)
add_code_block("# Create .env file\ncp .env.example .env\n\n# Edit .env with your settings:\n# DATABASE_URL=file:./db/custom.db  (SQLite for development)\n# JWT_SECRET=your-secret-key-change-in-production\n# NEXT_PUBLIC_APP_URL=http://localhost:3000")

add_para("Step 4: Initialize the database", bold=True)
add_code_block("# Generate Prisma client\nnpx prisma generate\n\n# Create database schema\nnpx prisma db push\n\n# Seed infrastructure data (88 sectors, 13 staff, 6 branches, etc.)\nnpx tsx scripts/seed.ts")

add_para("Step 5: Start the development server", bold=True)
add_code_block("# Using Bun (recommended)\nbun run dev\n\n# Using npm\nnpm run dev")

add_para("Step 6: Access the application", bold=True)
add_code_block("# Open in browser:\nhttp://localhost:3000\n\n# Login with credentials provided separately\n# by the system administrator.\n# Staff passwords are bcrypt-hashed and must\n# be communicated via secure channels only.")

add_heading("3.3 Seed Data", level=2)
add_para("The seed script (scripts/seed.ts) creates the following infrastructure data (NO demo customers or loans):")
add_table(
    ["Data", "Count", "Description"],
    [
        ["Branches", "6", "Lagos (2), Abuja (2), Ibadan, Benin"],
        ["Sectors", "88", "Business natures with benchmarked margins (Excel Sheet1 parity)"],
        ["Staff", "13", "One per role (super, md, cfo, hoc, cro, legal, 2 BMs, analyst, 2 LOs, frontdesk, treasury)"],
        ["Loan Products", "4", "Micro, SME, SME Plus, Exception"],
        ["Chart of Accounts", "30", "Standard GL accounts"],
        ["FAQ Articles", "15", "Customer knowledge base"],
    ],
    col_widths=[3.5, 2, 10.5]
)

# ============================================================================
# 4. DEPLOYMENT GUIDE
# ============================================================================
add_heading("4. Deployment Guide", level=1, page_break=True)

add_heading("4.1 Production Build", level=2)
add_code_block("# Build the production bundle\nbun run build\n\n# This creates:\n# - .next/standalone/ (self-contained server)\n# - .next/static/ (static assets)\n# - .next/standalone/.next/static/ (copied static)\n# - .next/standalone/public/ (copied public)")

add_heading("4.2 Start Production Server", level=2)
add_code_block("# Using the standalone server (recommended)\nNODE_ENV=production node .next/standalone/server.js\n\n# Or using Next.js start\nbun run start")

add_heading("4.3 Deployment Options", level=2)
add_heading("Option A: Vercel (Easiest)", level=3)
add_bullet("Push code to GitHub/GitLab")
add_bullet("Connect repository to Vercel")
add_bullet("Set environment variables in Vercel dashboard")
add_bullet("Deploy automatically on push to main")

add_heading("Option B: VPS / Dedicated Server", level=3)
add_code_block("# On the server:\ngit clone <repo>\ncd watershed-capital\nbun install\nbun run build\n\n# Use PM2 for process management\nnpm install -g pm2\npm2 start \"node .next/standalone/server.js\" --name watershed\npm2 save\npm2 startup")

add_heading("Option C: Docker", level=3)
add_code_block("# Dockerfile\nFROM node:20-alpine\nWORKDIR /app\nCOPY package*.json ./\nRUN npm install\nCOPY . .\nRUN npx prisma generate && npm run build\nEXPOSE 3000\nCMD [\"node\", \".next/standalone/server.js\"]")

add_heading("4.4 Caddy Reverse Proxy (Recommended)", level=2)
add_para("The project includes a Caddyfile for automatic HTTPS:")
add_code_block("# Caddyfile\nyourdomain.com {\n    reverse_proxy localhost:3000\n    encode gzip\n    header {\n        X-Frame-Options DENY\n        X-Content-Type-Options nosniff\n        Referrer-Policy strict-origin-when-cross-origin\n    }\n}")

# ============================================================================
# 5. ENVIRONMENT VARIABLES
# ============================================================================
add_heading("5. Environment Variables Template", level=1, page_break=True)

add_para("Create a .env file in the project root with the following variables:")
add_code_block("""# ============================================================================
# WATERSHED CAPITAL — ENVIRONMENT VARIABLES
# ============================================================================

# Database (SQLite for dev, PostgreSQL/Turso for production)
DATABASE_URL=file:./db/custom.db
# For production PostgreSQL:
# DATABASE_URL=postgresql://user:password@host:5432/watershed?schema=public

# JWT Authentication
JWT_SECRET=change-this-to-a-random-64-char-string-in-production
# Generate with: openssl rand -hex 32

# Application URL
NEXT_PUBLIC_APP_URL=http://localhost:3000

# Email Service (optional — system works without email)
# Supports: postmark, sendgrid, mailgun, nodemailer, aws-ses
EMAIL_PROVIDER=postmark
EMAIL_FROM=noreply@watershedcapital.com
EMAIL_API_KEY=your-email-provider-api-key

# SMS Service (optional)
SMS_PROVIDER=twilio
SMS_ACCOUNT_SID=your-twilio-sid
SMS_AUTH_TOKEN=your-twilio-token
SMS_FROM=+2348000000000

# CBN Regulatory Limits (configurable)
CBN_SINGLE_OBLIGOR_LIMIT=50000000
CBN_DSR_LIMIT=45
CBN_GEARING_LIMIT=35
CBN_COLLATERAL_MIN=100

# File Upload Path
UPLOAD_PATH=./uploads

# Node Environment
NODE_ENV=development""")

add_callout("Security Warning", "Never commit the .env file to version control. The JWT_SECRET must be changed from the default value before production deployment. Generate a secure secret with: openssl rand -hex 32")

# ============================================================================
# 6. DATABASE DOCUMENTATION
# ============================================================================
add_heading("6. Database Documentation (including ERD)", level=1, page_break=True)

add_heading("6.1 Database Overview", level=2)
add_para("The system uses Prisma ORM with 79 models spanning 2,240 lines of schema definition. The database is SQLite for development and can be migrated to PostgreSQL or Turso for production.")

add_heading("6.2 Core Models", level=2)
add_table(
    ["Model", "Purpose", "Key Relationships"],
    [
        ["Admin", "Staff accounts (13 roles)", "Belongs to Branch, has many Appraisals"],
        ["User", "Customer accounts", "Has one Business, has many Loans, Transactions"],
        ["Business", "Customer business info", "Belongs to User, has one Sector"],
        ["Branch", "Bank branches (6 seeded)", "Has many Staff, Customers, Loans"],
        ["Sector", "88 business sectors", "Has many Businesses, Loans"],
        ["LoanApplicants", "Loan applications", "Belongs to User, Branch, has one Appraisal"],
        ["CreditAppraisal", "CAM data (8 snapshots)", "Belongs to Loan, has many MCC Decisions"],
        ["MccDecision", "8-level approval decisions", "Belongs to Loan, Admin"],
        ["ApprovalLog", "Action audit trail", "Belongs to Loan, Admin"],
        ["LoanRepayment", "Repayment schedule", "Belongs to Loan"],
        ["LoanTransaction", "Disbursements & repayments", "Belongs to Loan"],
        ["ComplianceCondition", "Conditions precedent", "Belongs to Loan, has many Documents"],
        ["PreDisbursementChecklist", "22-item checklist", "Belongs to Loan (1:1)"],
        ["AuditLog", "System-wide audit", "Belongs to Admin"],
    ],
    col_widths=[4, 5, 7]
)

add_heading("6.3 CAM Snapshot Model (8-Layer Audit Trail)", level=2)
add_para("The CreditAppraisal model stores 8 frozen snapshots, one per approver role:")
add_table(
    ["Snapshot", "Field", "Created By", "When"],
    [
        ["LO Snapshot", "loSnapshot", "Loan Officer", "LO submits and locks CAM"],
        ["BM Snapshot", "bmSnapshot", "Branch Manager", "BM creates their assessment"],
        ["Analyst Snapshot", "analystSnapshot", "Credit Analyst", "Analyst completes structuring"],
        ["HOC Snapshot", "hocSnapshot", "Head of Credit", "HOC endorses structure"],
        ["CRO Snapshot", "croSnapshot", "CRO", "CRO completes risk assessment"],
        ["CFO Snapshot", "cfoSnapshot", "CFO", "CFO clears liquidity"],
        ["Legal Snapshot", "legalSnapshot", "Legal Officer", "Legal aggregates Executive Credit Pack"],
        ["MD Snapshot", "mdSnapshot", "MD/CEO", "MD gives final approval"],
    ],
    col_widths=[3.5, 3, 3.5, 6]
)

add_heading("6.4 Entity Relationship Summary", level=2)
add_code_block("""User (1) ──< (N) LoanApplicants (1) ── (1) CreditAppraisal
  │                    │                        │
  │                    │                        ├── (N) MccDecision
  │                    │                        ├── (N) ApprovalLog
  │                    │                        ├── (N) LoanDocument
  │                    │                        ├── (N) LoanRepayment
  │                    │                        ├── (N) LoanTransaction
  │                    │                        ├── (N) ComplianceCondition
  │                    │                        └── (1) PreDisbursementChecklist
  │                    │
  ├── (1) Business     ├── (N) BranchManagerVisit
  │         │          │
  │         └── Sector  └── (N) AuditLog
  │
  ├── (N) Transactions
  ├── (N) Savings
  ├── (N) TreasuryInvestment
  ├── (N) Ticket
  └── (N) Notification

Admin (1) ──< (N) LoanApplicants (as loanOfficer)
       │
       ├── (1) Branch (optional)
       ├── (N) AuditLog
       ├── (N) LoginHistory
       └── (N) ApprovalLog""")

# ============================================================================
# 7. API DOCUMENTATION
# ============================================================================
add_heading("7. API Documentation", level=1, page_break=True)

add_heading("7.1 Authentication", level=2)
add_para("All API routes (except /api/auth/login and /api/customer/login) require a JWT Bearer token in the Authorization header:")
add_code_block("Authorization: Bearer <jwt-token>")

add_para("Tokens are obtained by logging in:")
add_table(
    ["Endpoint", "Method", "Body", "Returns"],
    [
        ["/api/auth/login", "POST", "{ username, password }", "{ admin, token }"],
        ["/api/customer/login", "POST", "{ identifier, password }", "{ user, token }"],
    ],
    col_widths=[5, 2, 5, 4]
)

add_heading("7.2 Key API Endpoints", level=2)
add_table(
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/loans", "GET", "List loans (with filters: step, status, branch)"],
        ["/api/loans/[id]", "GET", "Get single loan with all relations"],
        ["/api/loans/[id]/transition", "POST", "Forward/return/query/reject/disburse loan"],
        ["/api/loans/[id]/disburse", "POST", "Disburse loan (creates repayment schedule)"],
        ["/api/loans/[id]/verify", "POST", "BVN/CAC verification"],
        ["/api/appraisals/[id]", "GET", "Load CAM data"],
        ["/api/appraisals/[id]", "PUT", "Save CAM data (with audit log)"],
        ["/api/engine/recalculate", "POST", "Run credit engine (30+ formulas)"],
        ["/api/mcc/[loanId]", "GET", "Get MCC decisions"],
        ["/api/mcc/[loanId]/decision", "POST", "Submit MCC decision"],
        ["/api/mcc/[loanId]/checklist", "GET/POST", "CP checklist management"],
        ["/api/mcc/[loanId]/export", "GET", "Export MCC paper (PDF)"],
        ["/api/sectors", "GET", "List 88 business sectors"],
        ["/api/branches", "GET", "List branches"],
        ["/api/staff", "GET/POST", "Staff management"],
        ["/api/customers?count=true", "GET", "Customer count (for SMS broadcast)"],
        ["/api/communications/announcements", "GET/POST", "Announcements"],
        ["/api/communications/messages", "GET", "Customer messages"],
        ["/api/communications/notifications-admin", "GET", "All notifications"],
        ["/api/dashboard/stats", "GET", "Dashboard statistics"],
        ["/api/audit/trail", "GET", "Audit trail"],
    ],
    col_widths=[6, 2, 8]
)

add_heading("7.3 Loan Transition API", level=2)
add_para("The most critical API — controls loan workflow transitions:")
add_code_block("""POST /api/loans/[id]/transition
Authorization: Bearer <token>
Content-Type: application/json

{
  "action": "forward" | "return" | "query" | "reject" | "disburse",
  "comment": "Optional comment",
  "nextStep": "Optional explicit next step",
  "mccDecision": {
    "recommendedAmount": 5000000,
    "duration": 12,
    "interestRatePercentage": 24,
    "ccdPercentage": 5,
    "upfrontFeePercentage": 3.2,
    "comment": "Decision comment",
    "decisionType": "approved" | "rejected" | "deferred" | "conditional"
  }
}

// Response:
{
  "action": "FORWARDED",
  "previousStep": "BM_QC",
  "newStep": "HOC_ASSIGNMENT",
  "newStatus": "processing"
}""")

# ============================================================================
# 8. ADMINISTRATOR MANUAL
# ============================================================================
add_heading("8. Administrator Manual", level=1, page_break=True)

add_heading("8.1 System Administrator Role", level=2)
add_para("The Super Admin (super.admin) has unrestricted access to all modules. Key responsibilities:")

add_bullet("Staff Management — Create/edit staff accounts, assign roles and branches")
add_bullet("Branch Management — Add/edit branches, assign branch managers")
add_bullet("Sector Management — Add/edit business sectors with benchmarked margins")
add_bullet("Loan Products — Configure loan products (Micro, SME, SME Plus, Exception)")
add_bullet("Global Settings — Configure site name, logo, currency, contact info")
add_bullet("Audit Trail — Monitor all system actions")
add_bullet("Branding — Customize logo, colors, tagline")

add_heading("8.2 Staff Roles & Permissions", level=2)
add_table(
    ["Role", "Scope", "Key Permissions"],
    [
        ["Super Admin", "All branches", "Full access (wildcard)"],
        ["MD/CEO", "All branches", "Final approval, general settings, portfolio"],
        ["CFO", "All branches", "Liquidity review, disbursement, treasury, accounting"],
        ["HOC", "All branches", "Structuring, assignment, finalization, supervisor"],
        ["CRO", "All branches", "Risk assessment, internal control"],
        ["Legal", "All branches", "KYC/CAC verification, aggregation, compliance"],
        ["BM", "Branch only", "Vetting, onboarding, KYC verify, branch manage"],
        ["Analyst", "All branches", "Credit analysis, structuring"],
        ["LO", "Branch only", "Origination, onboarding"],
        ["Front Desk", "Branch only", "Onboarding, support, messaging"],
        ["Treasury", "All branches", "Treasury onboarding, booking, assets"],
    ],
    col_widths=[3, 3, 10]
)

add_heading("8.3 Branch Scoping", level=2)
add_para("Branch-scoped roles (BM, LO, Front Desk, Treasury) can only access loans and customers assigned to their branch. National roles (Super, MD, CFO, HOC, CRO, Legal, Analyst) can access all branches. This is enforced at the API level in the transition route.")

add_heading("8.4 Daily Operations Checklist", level=2)
add_number("Check the Dashboard for pending loans at your gate")
add_number("Review and act on loans in your workflow queue")
add_number("Monitor audit trail for unusual activity")
add_number("Verify compliance conditions are being met")
add_number("Check NPL/Early Warning reports for portfolio health")

# ============================================================================
# 9. END-USER MANUAL
# ============================================================================
add_heading("9. End-User Manual", level=1, page_break=True)

add_heading("9.1 Staff Login", level=2)
add_number("Navigate to http://yourdomain.com (or localhost:3000)")
add_number("Click 'Staff Login'")
add_number("Enter your username and password")
add_number("Click 'Sign In'")
add_number("You will be redirected to the Dashboard")

add_heading("9.2 Loan Officer Workflow", level=2)
add_number("Click '1. LO Origination' in the sidebar")
add_number("Click 'New Application' to create a loan")
add_number("Fill in customer details, business info, loan amount")
add_number("Upload KYC documents (BVN, NIN, bank statement)")
add_number("Open the CAM (Credit Appraisal Memorandum)")
add_number("Fill all 12 CAM tabs (Profile, Business, Sales, Inventory, Expenses, Assets, Security, Visitation, Cross-Checks, Verifications, SWOT, Engine)")
add_number("Click 'Recalculate' to run the credit engine")
add_number("Fix any validation errors shown in the red panel")
add_number("Click 'Lock & Submit' to freeze the LO snapshot and forward to Legal")

add_heading("9.3 CAM (Credit Appraisal Memorandum)", level=2)
add_para("The Universal CAM has 12 tabs:")
add_table(
    ["Tab", "Purpose"],
    [
        ["1. Profile/KYC", "Customer identity, loan parameters, sector lookup, zonification"],
        ["2. Business", "Business description, financial inputs"],
        ["3. Sales Forensics", "4-source triangulation, stress test"],
        ["4. Stock Inventory", "Item-level margin analysis, weighted margin"],
        ["5. Expenses", "Business + family expenses (20% buffer)"],
        ["6. Assets/Balance Sheet", "Bank balances, fixed assets, liabilities, % of total"],
        ["7. Security & Guarantors", "Collateral registry, mix, guarantor info, biz verification"],
        ["8. Visitation", "LO + BM visitation reports (7 sections each), GPS, photo gallery"],
        ["9. Cross-Checks", "Zonification, loan cycle, capitalization, treasury, debt rotation, turnover"],
        ["10. Verifications", "Bank balances display, guarantor business verification"],
        ["11. SWOT", "Strengths, weaknesses, opportunities, threats, recommendation"],
        ["12. Engine Response", "Full engine output, 12-month projection, red flags, bank yield, detailed cashflow, BS comparison, amortization schedules"],
    ],
    col_widths=[4, 12]
)

add_heading("9.4 MCC Committee", level=2)
add_number("Navigate to MCC Committee → MCC Decisions")
add_number("Click on a loan to open the MCC Decision Paper")
add_number("Review the CAM, financial analysis, and all department recommendations")
add_number("Verify the 22-item CP Checklist (Conditions Precedent)")
add_number("Enter your decision (amount, tenor, rate, comment)")
add_number("Submit your decision")

add_heading("9.5 Customer Portal", level=2)
add_para("Customers access their portal at the same URL with their email/phone and password:")
add_bullet("Dashboard — Loan summary, active loans, notifications")
add_bullet("Apply for Loan — Submit new loan application")
add_bullet("My Loans — View loan details and repayment schedule")
add_bullet("Offers — View and accept/reject loan offers")
add_bullet("Pay Back — Make loan repayments")
add_bullet("Savings — View savings accounts")
add_bullet("Investments — View treasury investments")
add_bullet("Transactions — Transaction history")
add_bullet("Documents — Download loan documents")
add_bullet("Support — Submit support tickets, live chat, FAQ")

# ============================================================================
# 10. SECURITY & ACCESS CONTROL
# ============================================================================
add_heading("10. Security & Access Control Guide", level=1, page_break=True)

add_heading("10.1 Authentication", level=2)
add_bullet("JWT-based authentication using HMAC-SHA256 (Node.js crypto, no external dependency)")
add_bullet("Tokens expire after 8 hours")
add_bullet("Passwords are bcrypt-hashed (10 rounds)")
add_bullet("No password backdoors — bcrypt-only verification")
add_bullet("Generic error messages prevent user enumeration")

add_heading("10.2 API Security", level=2)
add_bullet("All API routes require Bearer token authentication")
add_bullet("Admin ID is extracted from JWT token, NOT request body (prevents impersonation)")
add_bullet("Branch scoping — branch roles can only access their branch's data")
add_bullet("Audit logging — every CAM edit, transition, and override is logged")

add_heading("10.3 Snapshot Locking", level=2)
add_para("Once a CAM is submitted (locked), the data becomes immutable:")
add_bullet("Only Super Admin or MD can override a lock")
add_bullet("Override requires a documented reason (minimum 10 characters)")
add_bullet("Override is audit-logged with 'critical' severity")
add_bullet("All fields changed during override are recorded")

add_heading("10.4 CBN Compliance Gates", level=2)
add_table(
    ["Gate", "Limit", "Action"],
    [
        ["DSR (Debt Service Ratio)", "≤ 45%", "Blocks submission if exceeded"],
        ["Gearing Ratio", "≤ 35%", "Blocks submission if exceeded"],
        ["Collateral Coverage", "≥ 100% FSV", "Critical red flag if below"],
        ["Single Obligor Limit", "₦50M (configurable)", "Blocks submission if exceeded"],
        ["Solvency Projection", "12-month positive", "Blocks submission if insolvent"],
        ["Engine Verdict", "APPROVE/REVIEW", "REJECT blocks submission"],
        ["AML CTR (Individual)", "₦5M", "Flag for Currency Transaction Report"],
        ["AML CTR (Corporate)", "₦10M", "Flag for Currency Transaction Report"],
    ],
    col_widths=[5, 4, 7]
)

add_heading("10.5 Session Management", level=2)
add_bullet("JWT token stored in localStorage")
add_bullet("Token cleared on logout")
add_bullet("Token cleared automatically on 401 response (expired/invalid)")
add_bullet("Zustand persist middleware stores admin session")

# ============================================================================
# 11. BACKUP & DISASTER RECOVERY
# ============================================================================
add_heading("11. Backup & Disaster Recovery Guide", level=1, page_break=True)

add_heading("11.1 Database Backup", level=2)
add_para("For SQLite (development):")
add_code_block("# Backup the database file\ncp db/custom.db db/backup/custom-$(date +%Y%m%d).db\n\n# Restore\ncp db/backup/custom-20250101.db db/custom.db")

add_para("For PostgreSQL (production):")
add_code_block("# Backup\npg_dump $DATABASE_URL > backup-$(date +%Y%m%d).sql\n\n# Restore\npsql $DATABASE_URL < backup-20250101.sql")

add_heading("11.2 Backup Schedule", level=2)
add_table(
    ["Frequency", "Type", "Retention"],
    [
        ["Daily", "Full database backup", "30 days"],
        ["Weekly", "Full database + code backup", "12 weeks"],
        ["Monthly", "Archival backup (offsite)", "12 months"],
        ["Before deploy", "Database snapshot", "Until next deploy"],
    ],
    col_widths=[3, 6, 7]
)

add_heading("11.3 Disaster Recovery", level=2)
add_number("Identify the issue (check server logs, database connectivity)")
add_number("If database corruption: restore from most recent backup")
add_number("If code issue: rollback to previous deployment")
add_number("If server failure: deploy to new server from Git repository")
add_number("Notify all stakeholders of downtime and recovery ETA")
add_number("Document the incident and root cause in the audit log")

add_heading("11.4 Automated Backup Script", level=2)
add_code_block("# Add to crontab for daily backups at 2 AM\n0 2 * * * cd /path/to/watershed && cp db/custom.db db/backup/custom-$(date +\\%Y\\%m\\%d).db && find db/backup/ -mtime +30 -delete")

# ============================================================================
# 12. MAINTENANCE & TROUBLESHOOTING
# ============================================================================
add_heading("12. Maintenance & Troubleshooting Guide", level=1, page_break=True)

add_heading("12.1 Routine Maintenance", level=2)
add_table(
    ["Task", "Frequency", "Command"],
    [
        ["Database backup", "Daily", "cp db/custom.db db/backup/"],
        ["Dependency update", "Monthly", "bun update && bun install"],
        ["Security audit", "Monthly", "bun audit"],
        ["Database integrity check", "Weekly", "npx prisma validate"],
        ["Log review", "Daily", "Check dev.log and server.log"],
        ["Disk space check", "Weekly", "df -h"],
    ],
    col_widths=[5, 3, 8]
)

add_heading("12.2 Common Issues & Solutions", level=2)
add_table(
    ["Issue", "Cause", "Solution"],
    [
        ["npm install ECONNRESET", "Network timeout on large packages", "Use: bun install"],
        ["Module not found: jsonwebtoken", "Next.js bundling Node.js module", "Fixed in v14 (uses crypto)"],
        ["Prisma client not generated", "Missing generate step", "Run: npx prisma generate"],
        ["Database locked", "Multiple processes accessing SQLite", "Restart the dev server"],
        ["401 Unauthorized on API", "Missing/expired JWT token", "Re-login to get new token"],
        ["CAM data not saving", "Locked snapshot", "Only super/MD can override with reason"],
        ["Duplicate key warning", "Same React key for multiple items", "Fixed in v14 (unique keys)"],
        ["Build fails with memory error", "Insufficient Node.js memory", "Set: NODE_OPTIONS=--max-old-space-size=4096"],
    ],
    col_widths=[4, 5, 7]
)

add_heading("12.3 Log Files", level=2)
add_bullet("dev.log — Development server output (auto-generated by bun run dev)")
add_bullet("server.log — Production server output (auto-generated by bun run start)")
add_bullet("Prisma query logs — Only in development (disabled in production)")

add_heading("12.4 Performance Optimization", level=2)
add_bullet("Use Bun instead of npm for 3x faster installs and builds")
add_bullet("Enable Turbopack (default in Next.js 16)")
add_bullet("Use standalone output for production (smaller deployment)")
add_bullet("Prisma query logging disabled in production")
add_bullet("Consider migrating from SQLite to PostgreSQL for high-traffic production")

# ============================================================================
# 13. SOURCE CODE REPOSITORY
# ============================================================================
add_heading("13. Source Code Repository Information", level=1, page_break=True)

add_heading("13.1 Repository Structure", level=2)
add_table(
    ["Directory", "Contents", "Description"],
    [
        ["src/app/api/", "137 route.ts files", "RESTful API endpoints"],
        ["src/app/page.tsx", "Main router", "Routes between admin/customer/public portals"],
        ["src/components/views/", "74 .tsx files", "UI views for all modules"],
        ["src/components/ui/", "shadcn components", "Reusable UI primitives"],
        ["src/components/pdf/", "4 PDF templates", "CAM memo, MCC paper, offer letter, receipt"],
        ["src/lib/", "15 .ts files", "Core libraries (engine, auth, db, constants)"],
        ["prisma/schema.prisma", "79 models", "Database schema definition"],
        ["scripts/seed.ts", "Seed script", "Infrastructure data (no demo data)"],
        ["public/", "Static assets", "Logo, images"],
    ],
    col_widths=[4, 4, 8]
)

add_heading("13.2 Key Source Files", level=2)
add_table(
    ["File", "Lines", "Purpose"],
    [
        ["src/components/views/cam.tsx", "4,200+", "Universal CAM with 12 tabs, 30+ financial features"],
        ["src/lib/credit-engine.ts", "2,600+", "Credit engine with 30+ formulas and CBN hard gates"],
        ["src/lib/constants.ts", "860+", "Workflow, roles, permissions, 88 sectors, 60+ locations"],
        ["src/components/views/mcc/mcc-detail.tsx", "1,610+", "MCC decision center with CP checklist"],
        ["src/components/views/loan/loan-detail.tsx", "890+", "Loan detail with workflow actions"],
        ["prisma/schema.prisma", "2,240+", "79 database models"],
        ["src/app/api/loans/[id]/transition/route.ts", "360+", "Loan workflow transition engine"],
        ["src/lib/auth.ts", "200+", "JWT authentication (HMAC-SHA256)"],
    ],
    col_widths=[7, 2, 7]
)

add_heading("13.3 Version Control", level=2)
add_bullet("Repository: Local development (configure Git remote for team collaboration)")
add_bullet("Branching strategy: main (production), develop (staging), feature/* (development)")
add_bullet("Commits should follow conventional commits format")
add_bullet("Current version: v14.0")

# ============================================================================
# 14. CREDENTIALS & ASSET HANDOVER CHECKLIST
# ============================================================================
add_heading("14. Credentials & Asset Handover Checklist", level=1, page_break=True)

add_heading("14.1 System Credentials", level=2)
add_callout("Security Notice", "All staff passwords are bcrypt-hashed (10 rounds) and stored securely in the database. Plaintext passwords are NEVER stored or documented. Initial passwords are set during the seeding process and MUST be changed immediately upon first login. Communicate passwords only via secure, encrypted channels (e.g., password manager, encrypted email, in-person handover).")
add_table(
    ["Role", "Username", "Password Status"],
    [
        ["Super Admin", "super.admin", "Must change on first login"],
        ["MD / CEO", "md", "Must change on first login"],
        ["CFO", "cfo", "Must change on first login"],
        ["HOC", "hoc", "Must change on first login"],
        ["CRO", "cro", "Must change on first login"],
        ["Legal Officer", "legal", "Must change on first login"],
        ["BM (Lagos)", "bm.lagos", "Must change on first login"],
        ["BM (Abuja)", "bm.abuja", "Must change on first login"],
        ["Credit Analyst", "analyst", "Must change on first login"],
        ["Loan Officer", "lo.lagos1", "Must change on first login"],
        ["Loan Officer", "lo.lagos2", "Must change on first login"],
        ["Front Desk", "frontdesk", "Must change on first login"],
        ["Treasury", "treasury", "Must change on first login"],
        ["JWT Secret", "(in .env file)", "Must be changed to random 64-char string"],
        ["Database URL", "(in .env file)", "Must be updated for production database"],
    ],
    col_widths=[4, 4, 8]
)

add_heading("14.2 Asset Handover Checklist", level=2)
checklist = [
    "Source code archive (watershed-capital-app-v14.zip)",
    "Database schema (prisma/schema.prisma)",
    "Seed script (scripts/seed.ts)",
    "Environment variables template (.env)",
    "Production build (.next/standalone/)",
    "SSL certificates (if applicable)",
    "Domain DNS access credentials",
    "Server SSH access credentials",
    "CI/CD pipeline configuration (if applicable)",
    "Email service API keys (Postmark/SendGrid)",
    "SMS service API keys (Twilio)",
    "Backup system configuration",
    "Monitoring/alerting setup",
    "User manual documentation (this document)",
    "Training session completed with operations team",
]
for item in checklist:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("☐  ")
    run.font.size = Pt(14)
    run2 = p.add_run(item)
    run2.font.size = Pt(11)

# ============================================================================
# 15. PROJECT ACCEPTANCE & SIGN-OFF
# ============================================================================
add_heading("15. Project Acceptance & Sign-off Document", level=1, page_break=True)

add_heading("15.1 Project Summary", level=2)
add_para("This document certifies that the Watershed Capital Banking Governance Platform has been developed, tested, and delivered in accordance with the agreed scope. The system implements a bank-grade 13-step loan origination workflow with CBN-compliant credit appraisal, multi-level approval chain, treasury, accounting, and post-disbursement monitoring.")

add_heading("15.2 Deliverables", level=2)
add_table(
    ["Deliverable", "Status", "Notes"],
    [
        ["13-step loan origination workflow", "✅ Complete", "Pre-Qualification → Structuring → Governance → Closing → Monitoring"],
        ["Universal CAM (12 tabs)", "✅ Complete", "Excel parity with 14 gap fixes"],
        ["Credit engine (30+ formulas)", "✅ Complete", "CBN hard gates: DSR 45%, Gearing 35%, Collateral 100%"],
        ["8-snapshot audit trail", "✅ Complete", "LO, BM, Analyst, HOC, CRO, CFO, Legal, MD"],
        ["MCC 8-level approval", "✅ Complete", "With 22-item CP checklist"],
        ["JWT authentication", "✅ Complete", "HMAC-SHA256, 8-hour expiry, no backdoors"],
        ["CAM PDF download", "✅ Complete", "CBN-standard memo format"],
        ["Auto-recalculation on MD approval", "✅ Complete", "Regenerates repayment schedule"],
        ["Customer negotiation state", "✅ Complete", "Dedicated workflow state for rejected offers"],
        ["Post-disbursement monitoring", "✅ Complete", "5 states: Active, Repayment, Warning, Collections, Closure"],
        ["Communication module", "✅ Complete", "7 features: Announcements, Messages, Notifications, Templates, SMS, Campaigns, Customer Service"],
        ["Treasury module", "✅ Complete", "Investors, products, booking, redemptions, assets, reports"],
        ["Accounting & GL", "✅ Complete", "Double-entry, journal, statements, reconciliation, invoicing, expenses, payroll, teller, tills, AP/AR"],
        ["Security hardening", "✅ Complete", "Branch scoping, audit logging, snapshot locking, input validation"],
        ["No demo/dummy data", "✅ Complete", "All INITIAL_DATA zeroed, DB verified clean"],
    ],
    col_widths=[6, 2.5, 7.5]
)

add_heading("15.3 Acceptance Criteria", level=2)
criteria = [
    "System builds without errors (next build exit code 0)",
    "All API routes require authentication (401 without token)",
    "Login works with bcrypt-verified passwords (no backdoors)",
    "CAM data persists and loads correctly (structured arrays mapped to JSON columns)",
    "CAM submission blocked if validation errors exist (20+ mandatory checks)",
    "Loan workflow transitions follow 13-step bank-grade flow",
    "MD approval triggers auto-recalculation of repayment schedule",
    "Locked snapshots cannot be edited (except by super/MD with reason)",
    "All CAM edits are audit-logged",
    "Branch scoping prevents cross-branch access",
    "No demo or dummy data in the database",
    "All 88 sectors seeded with benchmarked margins",
    "CBN hard gates enforced (DSR, Gearing, Collateral, Single Obligor)",
]
for c in criteria:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("✅  ")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_PRIMARY
    run2 = p.add_run(c)
    run2.font.size = Pt(11)

add_heading("15.4 Sign-off", level=2)
add_para("By signing below, the parties acknowledge that the project has been delivered as specified and all acceptance criteria have been met.")

# Sign-off table
sign_tbl = doc.add_table(rows=4, cols=3)
sign_tbl.style = 'Table Grid'
sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(["Role", "Name", "Signature & Date"]):
    cell = sign_tbl.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# Rows
sign_data = [
    ["Project Manager", "", ""],
    ["Technical Lead", "", ""],
    ["Client Representative", "", ""],
]
for r_idx, row in enumerate(sign_data):
    for c_idx, val in enumerate(row):
        cell = sign_tbl.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        # Add space for signature
        if c_idx == 2:
            p2 = cell.add_paragraph()
            p2.add_run("\n\n_______________________\nDate: ____________").font.size = Pt(9)

add_para("")
add_para("This document constitutes the complete project handover for the Watershed Capital Banking Governance Platform, Version 14.0.", italic=True, size=10, color=COLOR_GRAY)

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = "/home/z/my-project/download/Watershed-Capital-Project-Handover-Document.docx"
doc.save(output_path)
print(f"\n✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024:.1f} KB")
